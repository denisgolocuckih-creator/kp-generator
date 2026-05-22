import time
import streamlit as st
import base64
import os
import json
import tempfile
from datetime import datetime
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import pdfkit
import pandas as pd
from streamlit_mic_recorder import mic_recorder
from openai import OpenAI
import plotly.express as px
import plotly.graph_objects as go
import hashlib

# Инициализация темы
if "theme" not in st.session_state:
    st.session_state.theme = "light"

st.set_page_config(page_title="Генератор КП", page_icon="🏗️", layout="wide")

# ========== ПУТИ ==========
HISTORY_FILE = "data/history.json"
USERS_FILE = "data/users.json"

# ========== АВТОРИЗАЦИЯ ==========
def load_users():
    with open(USERS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def check_password(username, password):
    users = load_users()
    if username in users:
        return users[username]["password"] == password
    return False

def get_user_info(username):
    users = load_users()
    return users.get(username, None)

def init_session():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "username" not in st.session_state:
        st.session_state.username = ""
    if "user_info" not in st.session_state:
        st.session_state.user_info = None
    if "price_list" not in st.session_state:
        st.session_state.price_list = load_default_price_list()
    if "price_source" not in st.session_state:
        st.session_state.price_source = "Стандартный прайс-лист (Курск)"
    if "history" not in st.session_state:
        st.session_state.history = load_all_history() if st.session_state.get("authenticated") else []
    if "user_request" not in st.session_state:
        st.session_state.user_request = ""

# ========== ДАННЫЕ ==========
@st.cache_data
def load_default_price_list():
    with open("data/prices.txt", "r", encoding="utf-8") as f:
        return f.read()

@st.cache_data
def load_company_info():
    with open("data/company_info.txt", "r", encoding="utf-8") as f:
        return f.read()

def load_all_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_all_history(history):
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)

def get_user_history():
    all_history = load_all_history()
    if st.session_state.user_info and st.session_state.user_info["role"] == "admin":
        return all_history
    return [h for h in all_history if h.get("_user") == st.session_state.username]

def add_to_history(data):
    all_history = load_all_history()
    data["_user"] = st.session_state.username
    all_history.append(data)
    save_all_history(all_history)
    st.session_state.history = get_user_history()

def clear_all_history():
    save_all_history([])
    st.session_state.history = []

def clear_user_history():
    all_history = load_all_history()
    all_history = [h for h in all_history if h.get("_user") != st.session_state.username]
    save_all_history(all_history)
    st.session_state.history = []

company_info = load_company_info()
init_session()

# ========== СИСТЕМА ЛИЦЕНЗИЙ (ЗАЩИЩЁННАЯ) ==========
import hashlib
import uuid
import platform
import base64 as b64
from cryptography.fernet import Fernet

LICENSE_FILE = "data/license.key"
DEMO_FLAG_FILE = "data/.demo_used"

# Загружаем секреты
try:
    from secret import SECRET_SALT, ENCRYPTION_KEY
except ImportError:
    # Для облачной версии используем значения из secrets
    import os
    SECRET_SALT = os.getenv("SECRET_SALT", "CloudDefaultSalt2026!")
    ENCRYPTION_KEY = os.getenv("ENCRYPTION_KEY", "CloudEncryptionKey!")

# Инициализация шифрования
CIPHER = Fernet(b64.urlsafe_b64encode(hashlib.sha256(ENCRYPTION_KEY.encode()).digest()))

def get_hardware_id():
    """Получает уникальный идентификатор компьютера"""
    system = platform.system()
    node = platform.node()
    machine = platform.machine()
    raw = f"{system}-{node}-{machine}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16].upper()

def generate_license_key(days=30, hardware_id=None):
    """Генерирует лицензионный ключ с привязкой к железу"""
    expire_date = (datetime.now() + pd.Timedelta(days=days)).strftime('%Y.%m.%d')
    hw_id = hardware_id or get_hardware_id()
    raw = f"KP-{days}-{expire_date}-{hw_id}-{SECRET_SALT}"
    signature = hashlib.sha256(raw.encode()).hexdigest()[:16].upper()
    return f"KP-{days}-{expire_date}-{hw_id}-{signature}"

def check_license():
    """Проверяет лицензию с расшифровкой"""
    if not os.path.exists(LICENSE_FILE):
        return False, "Лицензия не найдена."
    
    try:
        # Читаем зашифрованный файл
        with open(LICENSE_FILE, "rb") as f:
            encrypted = f.read()
        
        # Расшифровываем
        decrypted = CIPHER.decrypt(encrypted).decode()
        key = decrypted.strip()
        
        parts = key.split("-")
        if len(parts) != 5 or parts[0] != "KP":
            return False, "Неверный формат ключа."
        
        days = int(parts[1])
        expire_date = parts[2]
        hw_id = parts[3]
        signature = parts[4]
        
        # Проверяем подпись
        raw = f"KP-{days}-{expire_date}-{hw_id}-{SECRET_SALT}"
        expected = hashlib.sha256(raw.encode()).hexdigest()[:16].upper()
        
        if signature != expected:
            return False, "Ключ недействителен (подпись)."
        
        # Проверяем привязку к железу
        current_hw = get_hardware_id()
        if hw_id != current_hw:
            return False, f"Ключ привязан к другому компьютеру (ожидается {hw_id}, текущий {current_hw})."
        
        # Проверяем срок
        expire = datetime.strptime(expire_date, '%Y.%m.%d')
        days_left = (expire - datetime.now()).days
        
        if days_left < 0:
            return False, f"Срок истёк ({expire_date})."
        
        return True, f"Лицензия активна. Осталось {days_left} дн."
    
    except Exception as e:
        return False, f"Ошибка лицензии: {str(e)[:50]}"

def save_license_key(key):
    """Сохраняет ключ в зашифрованном виде"""
    encrypted = CIPHER.encrypt(key.encode())
    with open(LICENSE_FILE, "wb") as f:
        f.write(encrypted)

def can_activate_demo():
    """Проверяет, можно ли активировать демо"""
    return not os.path.exists(DEMO_FLAG_FILE)

def mark_demo_used():
    """Помечает, что демо уже активировано"""
    with open(DEMO_FLAG_FILE, "w") as f:
        f.write(get_hardware_id())

# ========== СТРАНИЦА ВХОДА ==========
if not st.session_state.authenticated:
    st.title("🏗️ Генератор КП")
    st.markdown("### 🔐 Вход в систему")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        username = st.text_input("👤 Логин", placeholder="Введите логин")
        password = st.text_input("🔑 Пароль", type="password", placeholder="Введите пароль")
        
    if st.button("🚪 Войти", type="primary", use_container_width=True):
        if check_password(username, password):
                st.session_state.authenticated = True
                st.session_state.username = username
                st.session_state.user_info = get_user_info(username)
                st.session_state.history = get_user_history()
                # Проверяем лицензию при входе
                if os.path.exists(LICENSE_FILE):
                    valid, msg = check_license()
                    st.session_state.license_checked = valid
                    if not valid:
                        st.session_state.license_msg = msg
                else:
                    st.session_state.license_checked = False
                    st.session_state.license_msg = "Лицензия не найдена."
                st.rerun()
        else:
                st.error("❌ Неверный логин или пароль")
    
    st.divider()
    st.caption("💡 Тестовые пользователи:")
    st.caption("• admin / admin123 (видит все КП)")
    st.caption("• manager1 / manager123 (только свои КП)")
    st.caption("• manager2 / manager456 (только свои КП)")
    st.caption("• director / director789 (видит все КП)")
    st.stop()

# ========== ПРОВЕРКА ЛИЦЕНЗИИ ПОСЛЕ ВХОДА ==========
if st.session_state.get("authenticated") and not st.session_state.get("license_checked", True):
    st.title("🔑 Активация лицензии")
    st.markdown(f"### {st.session_state.get('license_msg', 'Лицензия не найдена.')}")
    st.caption(f"🖥️ ID компьютера: **{get_hardware_id()}**")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if can_activate_demo():
            st.markdown("### 🎁 Демо-доступ на 14 дней")
            st.caption("Демо-доступ можно активировать только один раз.")
            if st.button("🚀 Активировать демо", type="primary", use_container_width=True):
                demo_key = generate_license_key(14)
                save_license_key(demo_key)
                mark_demo_used()
                st.session_state.license_checked = True
                st.success("✅ Демо-доступ активирован на 14 дней!")
                st.rerun()
        else:
            st.warning("⚠️ Демо-доступ уже был использован.")
        
        st.divider()
        st.markdown("### 🔑 Введите лицензионный ключ")
        st.caption("Ключ привязан к ID вашего компьютера.")
        license_input = st.text_input("Ключ:", placeholder="KP-365-2026-12-31-ABCD1234-ABCD1234")
        if st.button("🔓 Активировать", use_container_width=True):
            if license_input:
                save_license_key(license_input)
                valid, msg = check_license()
                if valid:
                    st.session_state.license_checked = True
                    st.success("✅ Лицензия активирована!")
                    st.rerun()
                else:
                    st.error(f"❌ {msg}")
                    # Удаляем неверный файл
                    if os.path.exists(LICENSE_FILE):
                        os.remove(LICENSE_FILE)
    
    st.stop()

# Блокировка без лицензии
if st.session_state.get("authenticated") and not st.session_state.get("license_checked", True):
    st.warning("🔒 Программа не активирована. Введите лицензионный ключ или активируйте демо-доступ.")
    st.stop()

# ========== ИИ ==========
def get_llm():
    load_dotenv()
    return ChatOpenAI(
        model="deepseek/deepseek-chat",
        temperature=0.1,
        openai_api_key=os.getenv("VSEGPT_API_KEY"),
        openai_api_base="https://api.vsegpt.ru/v1"
    )

def calculate_kp(user_input):
    prompt = f"""
Ты — ИИ-калькулятор строительной компании. Твоя задача: выдать ТОЧНЫЙ математический расчёт.

=== БАЗА ЗНАНИЙ (ЦЕНЫ И КОЭФФИЦИЕНТЫ) ===
{st.session_state.price_list}
=== КОНЕЦ БАЗЫ ===

=== ДАННЫЕ КОМПАНИИ ===
{company_info}
=== КОНЕЦ ДАННЫХ ===

АЛГОРИТМ РАСЧЁТА (ВЫПОЛНЯЙ ПО ШАГАМ):
Шаг 1. Найди в базе точную базовую ставку для указанного типа объекта и класса.
Шаг 2. Базовая стоимость = Площадь × Базовая ставка.
Шаг 3. Коэффициенты удорожания применяй ПОСЛЕДОВАТЕЛЬНО (каждый к базе, итог нарастающим).
Шаг 4. Скидки применяй ПОСЛЕДОВАТЕЛЬНО к итогу после коэффициентов.
Шаг 5. Срок строительства = Площадь / 50. Округли до целого вверх.

ФОРМАТ ОТВЕТА — ТОЛЬКО JSON:
{{
  "тип_объекта": "...",
  "площадь_квм": ...,
  "базовая_ставка_за_квм": ...,
  "базовая_стоимость": ...,
  "коэффициенты": [{{"название": "...", "процент": ..., "сумма": ...}}],
  "стоимость_после_коэффициентов": ...,
  "скидки": [{{"название": "...", "процент": ..., "сумма": ...}}],
  "итоговая_стоимость": ...,
  "срок_месяцев": ...,
  "комментарий": "..."
}}
"""
    result = get_llm().invoke(f"{prompt}\n\nЗАПРОС: {user_input}\n\nВерни ТОЛЬКО JSON:")
    return result.content

def generate_word(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    company = st.session_state.get('company_name', 'Моя компания')
    phone = st.session_state.get('company_phone', '+7 (4712) 123-45-67')
    email = st.session_state.get('company_email', 'info@stroyinvest46.ru')
    inn = st.session_state.get('company_inn', '4632123456')
    
    header = doc.sections[0].header
    h_para = header.paragraphs[0]
    if not st.session_state.logo.startswith("http"):
        logo_bytes = base64.b64decode(st.session_state.logo)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(logo_bytes)
            tmp_path = tmp.name
        run = h_para.add_run()
        run.add_picture(tmp_path, width=Inches(0.8))
        os.remove(tmp_path)
    h_para.add_run(f"  ООО «{company}» | Курск | {phone} | {email}")
    h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    title = doc.add_heading('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}").font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_heading('Объект строительства', level=2)
    info_table = doc.add_table(rows=5, cols=2, style='Light List Accent 1')
    info_table.cell(0, 0).text = "Тип объекта:"
    info_table.cell(0, 1).text = str(data.get("тип_объекта", "—"))
    info_table.cell(1, 0).text = "Площадь:"
    info_table.cell(1, 1).text = f"{data.get('площадь_квм', 0):,} м²".replace(",", " ")
    info_table.cell(2, 0).text = "Срок строительства:"
    info_table.cell(2, 1).text = f"~{data.get('срок_месяцев', '—')} мес."
    info_table.cell(3, 0).text = "Город:"
    info_table.cell(3, 1).text = "Курск"
    info_table.cell(4, 0).text = "Менеджер:"
    info_table.cell(4, 1).text = st.session_state.user_info["name"] if st.session_state.user_info else "—"
    
    doc.add_paragraph()
    doc.add_heading('Расчёт стоимости', level=2)
    cost_table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
    for cell, text in zip(cost_table.rows[0].cells, ["Статья", "Ставка / %", "Сумма, ₽", "Итого, ₽"]):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    
    base_rate = data.get("базовая_ставка_за_квм", 0)
    base_cost = data.get("базовая_стоимость", 0)
    
    row = cost_table.add_row().cells
    row[0].text = f"Базовая стоимость ({data.get('площадь_квм', 0)} м² × {base_rate:,} ₽/м²)".replace(",", " ")
    row[1].text = "—"
    row[2].text = f"{base_cost:,.0f}".replace(",", " ")
    row[3].text = f"{base_cost:,.0f}".replace(",", " ")
    
    current = base_cost
    for c in data.get("коэффициенты", []):
        current += c["сумма"]
        row = cost_table.add_row().cells
        row[0].text = c["название"]
        row[1].text = f"+{c['процент']}%"
        row[2].text = f"{c['сумма']:,.0f}".replace(",", " ")
        row[3].text = f"{current:,.0f}".replace(",", " ")
    
    for d in data.get("скидки", []):
        current -= d["сумма"]
        row = cost_table.add_row().cells
        row[0].text = d["название"]
        row[1].text = f"-{d['процент']}%"
        row[2].text = f"-{d['сумма']:,.0f}".replace(",", " ")
        row[3].text = f"{current:,.0f}".replace(",", " ")
    
    row = cost_table.add_row().cells
    row[0].text = "ИТОГО"
    row[3].text = f"{data.get('итоговая_стоимость', 0):,.0f}".replace(",", " ")
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    
    doc.add_paragraph()
    if data.get("комментарий"):
        doc.add_heading('Примечания', level=2)
        doc.add_paragraph(data["комментарий"])
    
    doc.add_paragraph()
    doc.add_heading('Реквизиты исполнителя', level=2)
    doc.add_paragraph(f"ООО «{company}»")
    doc.add_paragraph(f"ИНН: {inn} | КПП: 463201001")
    doc.add_paragraph("305000, г. Курск, ул. Ленина, д. 15, офис 301")
    doc.add_paragraph(f"Тел.: {phone} | Email: {email}")
    doc.add_paragraph()
    doc.add_paragraph(f"Менеджер: {st.session_state.user_info['name'] if st.session_state.user_info else '—'}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_excel(history):
    if not history:
        return None
    rows = []
    for h in history:
        coeffs = ", ".join([f"{c['название']} (+{c['процент']}%)" for c in h.get("коэффициенты", [])])
        discounts = ", ".join([f"{d['название']} (-{d['процент']}%)" for d in h.get("скидки", [])])
        rows.append({
            "Менеджер": h.get("_user", ""),
            "Дата": h.get("_время", ""),
            "Тип объекта": h.get("тип_объекта", ""),
            "Площадь, м²": h.get("площадь_квм", 0),
            "Базовая ставка, ₽/м²": h.get("базовая_ставка_за_квм", 0),
            "Базовая стоимость, ₽": h.get("базовая_стоимость", 0),
            "Коэффициенты": coeffs,
            "Итоговая стоимость, ₽": h.get("итоговая_стоимость", 0),
            "Срок, мес.": h.get("срок_месяцев", 0),
        })
    df = pd.DataFrame(rows)
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="Все КП")
        ws = writer.sheets["Все КП"]
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)
    buffer.seek(0)
    return buffer

def generate_pdf(data):
    config = pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')
    
    company = st.session_state.get('company_name', 'Моя компания')
    phone = st.session_state.get('company_phone', '+7 (4712) 123-45-67')
    
    html = f"""
    <html><head><meta charset="utf-8"><style>
        body {{ font-family: Arial; margin: 40px; }}
        h1 {{ text-align: center; }}
        table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
        td, th {{ border: 1px solid #ddd; padding: 8px; }}
        th {{ background: #f5f5f5; }}
        .header {{ text-align: center; color: gray; font-size: 10px; }}
        .sign {{ margin-top: 50px; text-align: right; }}
    </style></head><body>
        <div class="header">ООО «{company}» | Курск | {phone}</div>
        <h1>КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ</h1>
        <p>Дата: {datetime.now().strftime('%d.%m.%Y')}</p>
        <table>
            <tr><td>Тип объекта:</td><td>{data.get('тип_объекта', '—')}</td></tr>
            <tr><td>Площадь:</td><td>{data.get('площадь_квм', 0):,} м²</td></tr>
            <tr><td>Срок:</td><td>~{data.get('срок_месяцев', '—')} мес.</td></tr>
        </table>
        <h2>Расчёт</h2>
        <table>
            <tr><th>Статья</th><th>%</th><th>Сумма, ₽</th><th>Итого, ₽</th></tr>
            <tr><td>Базовая стоимость</td><td>—</td><td>{data.get('базовая_стоимость', 0):,.0f}</td><td>{data.get('базовая_стоимость', 0):,.0f}</td></tr>
    """
    current = data.get('базовая_стоимость', 0)
    for c in data.get('коэффициенты', []):
        current += c['сумма']
        html += f"<tr><td>{c['название']}</td><td>+{c['процент']}%</td><td>{c['сумма']:,.0f}</td><td>{current:,.0f}</td></tr>"
    for d in data.get('скидки', []):
        current -= d['сумма']
        html += f"<tr><td>{d['название']}</td><td>-{d['процент']}%</td><td>-{d['сумма']:,.0f}</td><td>{current:,.0f}</td></tr>"
    html += f"<tr><td colspan='3'><b>ИТОГО</b></td><td><b>{data.get('итоговая_стоимость', 0):,.0f}</b></td></tr></table>"
    html += f"<p>ООО «{company}», Курск</p>"
    html += f"""
        <div class="sign">
            <p>_________________________ / Иванов С.П.</p>
            <p style="font-size: 10px; color: gray;">Генеральный директор ООО «{company}»</p>
            <p style="font-size: 10px; color: gray;">М.П.</p>
        </div>
    </body></html>
    """
    
    return pdfkit.from_string(html, False, configuration=config)

def generate_xml(data):
    company = st.session_state.get('company_name', 'Моя компания')
    inn = st.session_state.get('company_inn', '4632123456')
    
    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<КоммерческоеПредложение>
    <Номер>КП-{datetime.now().strftime('%Y%m%d-%H%M')}</Номер>
    <Дата>{datetime.now().strftime('%Y.%m.%d')}</Дата>
    <Поставщик>
        <Наименование>ООО «{company}»</Наименование>
        <ИНН>{inn}</ИНН>
    </Поставщик>
    <Объект>
        <Тип>{data.get('тип_объекта', '')}</Тип>
        <Площадь>{data.get('площадь_квм', 0)}</Площадь>
        <СрокСтроительстваМесяцев>{data.get('срок_месяцев', 0)}</СрокСтроительстваМесяцев>
    </Объект>
    <Расчёт>
        <БазоваяСтавкаЗаКвМ>{data.get('базовая_ставка_за_квм', 0)}</БазоваяСтавкаЗаКвМ>
        <БазоваяСтоимость>{data.get('базовая_стоимость', 0)}</БазоваяСтоимость>
        <СтоимостьПолеКоэффициентов>{data.get('стоимость_после_коэффициентов', 0)}</СтоимостьПолеКоэффициентов>
        <ИтоговаяСтоимость>{data.get('итоговая_стоимость', 0)}</ИтоговаяСтоимость>
    </Расчёт>
    <Коэффициенты>
"""
    for c in data.get('коэффициенты', []):
        xml += f"""        <Коэффициент>
            <Название>{c['название']}</Название>
            <Процент>{c['процент']}</Процент>
            <Сумма>{c['сумма']}</Сумма>
        </Коэффициент>
"""
    xml += """    </Коэффициенты>
    <Скидки>
"""
    for d in data.get('скидки', []):
        xml += f"""        <Скидка>
            <Название>{d['название']}</Название>
            <Процент>{d['процент']}</Процент>
            <Сумма>{d['сумма']}</Сумма>
        </Скидка>
"""
    xml += """    </Скидки>
    <Комментарий>"""
    xml += data.get('комментарий', '')
    xml += """</Комментарий>
</КоммерческоеПредложение>"""
    
    return xml.encode('utf-8')

# ========== БОКОВАЯ ПАНЕЛЬ ==========
with st.sidebar:
    if "logo" not in st.session_state:
        st.session_state.logo = "https://img.icons8.com/fluency/96/building.png"
    
    uploaded_logo = st.file_uploader("📷 Загрузите логотип (PNG, JPG):", type=["png", "jpg", "jpeg"], key="logo_uploader")
    if uploaded_logo is not None:
        st.session_state.logo = base64.b64encode(uploaded_logo.read()).decode()
        st.success("✅ Логотип загружен!")
    
    st.image(st.session_state.logo if st.session_state.logo.startswith("http") else f"data:image/png;base64,{st.session_state.logo}", width=200)
    st.title(f"🏗️ {st.session_state.get('company_name', 'Моя компания')}")
    
    user_name = st.session_state.user_info["name"] if st.session_state.user_info else "Гость"
    user_role = st.session_state.user_info["role"] if st.session_state.user_info else ""
    st.markdown(f"👤 **{user_name}**")
    st.caption(f"Роль: {user_role} | Логин: {st.session_state.username}")
    
    if st.button("🚪 Выйти", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()
    
    st.markdown("**ИИ-генератор КП v8.0**")
    if st.session_state.get("authenticated"):
        valid, msg = check_license()
        if valid:
            st.caption(f"🔒 {msg}")
        else:
            st.caption("⚠️ Требуется активация")
    st.divider()
    
    theme_toggle = st.toggle("🌙 Тёмная тема", value=(st.session_state.theme == "dark"))
    st.session_state.theme = "dark" if theme_toggle else "light"
    
    if st.session_state.theme == "dark":
        st.markdown("""
        <style>
            .stApp { background-color: #1a1a2e; color: #e0e0e0; }
            .stButton>button { background-color: #16213e; color: white; }
            .stTextArea>div>div>textarea { background-color: #16213e; color: white; }
            section[data-testid="stSidebar"] { background-color: #0f0f23; }
        </style>
        """, unsafe_allow_html=True)
    
    st.divider()
    
    st.markdown("### ⚙️ Настройки компании")
    if "company_name" not in st.session_state:
        st.session_state.company_name = "Моя компания"
    if "company_phone" not in st.session_state:
        st.session_state.company_phone = "+7 (4712) 123-45-67"
    if "company_email" not in st.session_state:
        st.session_state.company_email = "info@example.ru"
    if "company_inn" not in st.session_state:
        st.session_state.company_inn = "4632123456"
    
    with st.expander("⚙️ Изменить реквизиты"):
        st.session_state.company_name = st.text_input("Название компании:", value=st.session_state.company_name)
        st.session_state.company_phone = st.text_input("Телефон:", value=st.session_state.company_phone)
        st.session_state.company_email = st.text_input("Email:", value=st.session_state.company_email)
        st.session_state.company_inn = st.text_input("ИНН:", value=st.session_state.company_inn)
        if st.button("💾 Сохранить реквизиты", use_container_width=True):
            st.success("✅ Реквизиты обновлены!")
    
    st.divider()
    
    st.markdown("### 📊 Прайс-лист")
    st.caption(f"Активен: **{st.session_state.price_source}**")
    
    uploaded_file = st.file_uploader("Загрузите прайс-лист:", type=["xlsx", "csv", "txt"])
    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith(".txt"):
                st.session_state.price_list = uploaded_file.read().decode("utf-8")
            elif uploaded_file.name.endswith(".csv"):
                st.session_state.price_list = pd.read_csv(uploaded_file).to_string(index=False)
            elif uploaded_file.name.endswith(".xlsx"):
                st.session_state.price_list = pd.read_excel(uploaded_file).to_string(index=False)
            st.session_state.price_source = f"📁 {uploaded_file.name}"
            st.success(f"✅ Загружен")
        except Exception as e:
            st.error(f"Ошибка: {e}")
    
    if st.button("🔄 Сбросить", use_container_width=True):
        st.session_state.price_list = load_default_price_list()
        st.session_state.price_source = "Стандартный"
        st.rerun()
    
    st.divider()
    
    st.markdown("### 📋 Шаблоны КП")
    if "templates" not in st.session_state:
        st.session_state.templates = {}
    
    template_name = st.text_input("Название шаблона:", placeholder="Например: Эконом-офис", key="tpl_name")
    template_text = st.text_area("Текст шаблона:", placeholder="Офис 500 кв.м., эконом, без отделки...", height=80, key="tpl_text")
    
    col_t1, col_t2 = st.columns(2)
    with col_t1:
        if st.button("💾 Сохранить", use_container_width=True, key="save_tpl"):
            if template_name and template_text:
                st.session_state.templates[template_name] = template_text
                st.success(f"✅ Шаблон '{template_name}' сохранён!")
    with col_t2:
        if st.button("📂 Загрузить", use_container_width=True, key="load_tpl"):
            if st.session_state.templates:
                tpl_names = list(st.session_state.templates.keys())
                selected = st.selectbox("Выберите:", tpl_names, key="sel_tpl")
                if selected:
                    st.session_state.user_request = st.session_state.templates[selected]
                    st.rerun()
    
    st.divider()
    st.markdown("### 📋 Примеры:")
    for ex in ["Офис 500 кв.м., бизнес-класс, центр, срочно", "Склад 2000 кв.м., тёплый, эконом", "Коттедж 350 кв.м., элитный, зима"]:
        if st.button(ex, use_container_width=True):
            st.session_state.user_request = ex
    
    st.divider()
    if st.session_state.history:
        excel_buffer = generate_excel(st.session_state.history)
        st.download_button("📥 Excel (все КП)", excel_buffer, f"КП_{datetime.now().strftime('%Y%m%d')}.xlsx", use_container_width=True)
    
    st.divider()
    st.markdown("### 📦 Пакетная генерация")
    batch_file = st.file_uploader("Excel со списком объектов:", type=["xlsx"], key="batch_upload")
    if batch_file is not None:
        try:
            df_batch = pd.read_excel(batch_file)
            st.caption(f"Найдено объектов: **{len(df_batch)}**")
            if st.button(f"⚡ Сгенерировать {len(df_batch)} КП", use_container_width=True, type="primary"):
                progress_bar = st.progress(0, text="Начинаем генерацию...")
                for i, (_, row) in enumerate(df_batch.iterrows()):
                    desc = f"{row.get('Тип', '')} {row.get('Площадь', '')} кв.м., {row.get('Класс', '')}, {row.get('Условия', '')}"
                    progress_bar.progress((i) / len(df_batch), text=f"📝 {desc[:60]}...")
                    try:
                        resp = calculate_kp(desc)
                        s, e = resp.find('{'), resp.rfind('}') + 1
                        if s != -1 and e > s:
                            data = json.loads(resp[s:e])
                            data["_время"] = datetime.now().strftime(" %Y.%m.%dH:%M")
                            data["_статус"] = "Черновик"
                            data["_user"] = st.session_state.username
                            add_to_history(data)
                    except:
                        pass
                progress_bar.progress(100, text="✅ Готово!")
                st.success(f"✅ Сгенерировано {len(df_batch)} КП!")
                st.rerun()
        except Exception as e:
            st.error(f"Ошибка чтения файла: {e}")
    
    st.divider()
    st.caption(f"© 2026 {st.session_state.get('company_name', 'Моя компания')} | Сохранено: {len(st.session_state.history)} КП")

# ========== ВКЛАДКИ ==========
if st.session_state.get("authenticated") and not st.session_state.get("license_checked", True):
    st.warning("🔒 Программа не активирована.")
    st.stop()

tab_main, tab_analytics, tab_chat = st.tabs(["🧮 Калькулятор КП", "📊 Аналитика", "💬 Чат с ИИ"])

with tab_main:
    st.title("🧠 Генератор коммерческих предложений")
    st.markdown("Опишите объект или **надиктуйте голосом**.")
    
    st.markdown("### 🎤 Голосовой ввод")
    audio = mic_recorder(start_prompt="🎤 Начать запись", stop_prompt="⏹️ Остановить", just_once=True, use_container_width=True, key="mic")
    
    if audio and audio.get("bytes"):
        with st.spinner("🎙️ Распознаю..."):
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
                f.write(audio["bytes"])
                temp_path = f.name
            load_dotenv()
            client = OpenAI(api_key=os.getenv("VSEGPT_API_KEY"), base_url="https://api.vsegpt.ru/v1")
            with open(temp_path, "rb") as af:
                transcript = client.audio.transcriptions.create(model="stt-openai/whisper-1", file=af, language="ru")
            os.remove(temp_path)
            st.session_state.user_request = transcript.text
            st.rerun()
    
    user_request = st.text_area("📝 Описание объекта:", value=st.session_state.user_request, placeholder="Офис 500 кв.м., бизнес-класс, центр...", height=100)
    
    c1, c2, c3, c4, c5 = st.columns([2, 1, 2, 1, 2])
    with c3:
        btn = st.button("🚀 Рассчитать КП", type="primary", use_container_width=True)
    with c5:
        if st.button("🗑️ Очистить", use_container_width=True):
            st.session_state.user_request = ""
            st.rerun()
    
    if btn and user_request:
        with st.spinner(""):
            progress_bar = st.progress(0, text="🔍 Анализ запроса...")
            time.sleep(0.3)
            progress_bar.progress(25, text="📊 Поиск ставок...")
            time.sleep(0.3)
            progress_bar.progress(50, text="🧮 Расчёт стоимости...")
            time.sleep(0.3)
            progress_bar.progress(75, text="📝 Формирование документа...")
            
            try:
                resp = calculate_kp(user_request)
                progress_bar.progress(100, text="✅ Готово!")
                time.sleep(0.3)
                progress_bar.empty()
                
                s, e = resp.find('{'), resp.rfind('}') + 1
                if s != -1 and e > s:
                    data = json.loads(resp[s:e])
                    data["_время"] = datetime.now().strftime("%Y.%m.%d %H:%M")
                    data["_статус"] = "Черновик"
                    add_to_history(data)
                    st.success("✅ Готово!")
                    
                    c1, c2, c3 = st.columns(3)
                    c1.metric("🏢 Тип", data.get("тип_объекта", "—"))
                    c2.metric("📐 Площадь", f"{data.get('площадь_квм', 0):,} м²".replace(",", " "))
                    c3.metric("💰 Итого", f"{data.get('итоговая_стоимость', 0):,.0f} ₽".replace(",", " "))
                    
                    st.divider()
                    t1, t2, t3, t4 = st.tabs(["📋 Сводка", "📈 Коэфф.", "🏷️ Скидки", "🏢 Контакты"])
                    with t1:
                        br = data.get("базовая_ставка_за_квм", 0)
                        bc = data.get("базовая_стоимость", 0)
                        af = data.get("стоимость_после_коэффициентов", "—")
                        ca, cb = st.columns(2)
                        ca.metric("Базовая ставка", f"{br:,} ₽/м²".replace(",", " "))
                        ca.metric("Базовая стоимость", f"{bc:,.0f} ₽".replace(",", " "))
                        cb.metric("После коэфф.", f"{af:,.0f} ₽".replace(",", " ") if isinstance(af, (int, float)) else af)
                        cb.metric("Итого", f"{data.get('итоговая_стоимость', 0):,.0f} ₽".replace(",", " "))
                        st.metric("📅 Срок", f"~{data.get('срок_месяцев', '—')} мес.")
                        st.info(data.get("комментарий", "—"))
                    with t2:
                        for c in data.get("коэффициенты", []):
                            st.write(f"• {c['название']}: +{c['процент']}% (+{c['сумма']:,.0f} ₽)".replace(",", " "))
                    with t3:
                        for d in data.get("скидки", []):
                            st.write(f"• {d['название']}: -{d['процент']}% (-{d['сумма']:,.0f} ₽)".replace(",", " "))
                    with t4:
                        st.write(f"ООО «{st.session_state.get('company_name', 'Моя компания')}», Курск")
                    
                    st.divider()
                    dc1, dc2, dc3, dc4 = st.columns(4)
                    dc1.download_button("📥 JSON", json.dumps(data, ensure_ascii=False, indent=2), f"KP_{datetime.now().strftime('%H%M')}.json", use_container_width=True)
                    dc2.download_button("📄 Word", generate_word(data), f"KP_{datetime.now().strftime('%H%M')}.docx", use_container_width=True)
                    dc3.download_button("📕 PDF", generate_pdf(data), f"KP_{datetime.now().strftime('%H%M')}.pdf", "application/pdf", use_container_width=True)
                    dc4.download_button("📋 XML (1С)", generate_xml(data), f"KP_{datetime.now().strftime('%H%M')}.xml", "application/xml", use_container_width=True)
                    
                    dc4_col, dc5_col = st.columns(2)
                    with dc4_col:
                        kp_json = json.dumps(data, ensure_ascii=False)
                        kp_b64 = base64.b64encode(kp_json.encode()).decode()
                        share_url = f"http://localhost:8501/?kp={kp_b64}"
                        st.text_input("🔗 Ссылка на КП:", value=share_url, key="share_url")
                        st.caption("Скопируйте ссылку, чтобы поделиться КП")
                    
                    st.divider()
                    st.markdown("### ✏️ Статус КП")
                    kp_list = [f"{h.get('_время', '')} | {h.get('тип_объекта', '')}" for h in st.session_state.history]
                    if kp_list:
                        col_s1, col_s2, col_s3 = st.columns([2, 2, 1])
                        with col_s1:
                            selected_kp = st.selectbox("Выберите КП:", range(len(kp_list)), format_func=lambda x: kp_list[x], key="sel_kp_status_1")
                        with col_s2:
                            new_status = st.selectbox("Новый статус:", ["Черновик", "Отправлено", "Согласовано", "Отказ"], key="sel_status_1")
                        with col_s3:
                            if st.button("✏️ Обновить", use_container_width=True):
                                st.session_state.history[selected_kp]["_статус"] = new_status
                                save_all_history(st.session_state.history)
                                st.success(f"Статус изменён на '{new_status}'")
                                st.rerun()
            except Exception as ex:
                st.error(f"Ошибка: {ex}")
    
    st.divider()
    st.caption("💡 Голосовой ввод — в Chrome.")

# ========== АНАЛИТИКА ==========
with tab_analytics:
    st.title("📊 Аналитика и сравнение КП")
    
    if not st.session_state.history:
        st.info("👆 Сделайте несколько расчётов во вкладке «Калькулятор».")
    else:
        df = pd.DataFrame(st.session_state.history)
        
        subtab1, subtab2 = st.tabs(["📈 Общая аналитика", "🔄 Сравнение КП"])
        
        with subtab1:
            st.markdown("### 🔍 Фильтры")
            col_f1, col_f2, col_f3 = st.columns(3)
            
            with col_f1:
                if st.session_state.user_info and st.session_state.user_info["role"] == "admin":
                    users_list = ["Все"] + list(df["_user"].unique()) if "_user" in df.columns else ["Все"]
                    selected_user = st.selectbox("👤 Менеджер:", users_list)
                    if selected_user != "Все":
                        df = df[df["_user"] == selected_user]
            
            with col_f2:
                if "тип_объекта" in df.columns:
                    types_list = ["Все"] + list(df["тип_объекта"].unique())
                    selected_type = st.selectbox("🏢 Тип объекта:", types_list)
                    if selected_type != "Все":
                        df = df[df["тип_объекта"] == selected_type]
            
            with col_f3:
                date_range = st.selectbox("📅 Период:", ["Всё время", "Сегодня", "Последние 7 дней", "Последние 30 дней"])
                today = datetime.now().date()
                if date_range == "Сегодня":
                    df = df[df["_время"].str[:10] == str(today)]
                elif date_range == "Последние 7 дней":
                    df = df[df["_время"].str[:10] >= str(today - pd.Timedelta(days=7))]
                elif date_range == "Последние 30 дней":
                    df = df[df["_время"].str[:10] >= str(today - pd.Timedelta(days=30))]
            
            if len(df) == 0:
                st.warning("Нет данных по выбранным фильтрам.")
            else:
                st.divider()
                
                total_kp = len(df)
                total_sum = df["итоговая_стоимость"].sum()
                avg_check = df["итоговая_стоимость"].mean()
                avg_area = df["площадь_квм"].mean()
                
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("📋 КП", total_kp)
                m2.metric("💰 Сумма", f"{total_sum:,.0f} ₽".replace(",", " "))
                m3.metric("💵 Средний чек", f"{avg_check:,.0f} ₽".replace(",", " "))
                m4.metric("📐 Средняя площадь", f"{avg_area:,.0f} м²".replace(",", " "))
                
                if "_статус" in df.columns:
                    st.divider()
                    st.markdown("### 📊 Воронка продаж")
                    statuses = ["Черновик", "Отправлено", "Согласовано", "Отказ"]
                    status_counts = {s: len(df[df["_статус"] == s]) for s in statuses}
                    
                    col_s1, col_s2, col_s3, col_s4 = st.columns(4)
                    with col_s1:
                        st.metric("📝 Черновик", status_counts["Черновик"])
                    with col_s2:
                        st.metric("📤 Отправлено", status_counts["Отправлено"])
                    with col_s3:
                        st.metric("✅ Согласовано", status_counts["Согласовано"])
                    with col_s4:
                        st.metric("❌ Отказ", status_counts["Отказ"])
                    
                    total_with_status = status_counts["Черновик"] + status_counts["Отправлено"]
                    if total_with_status > 0:
                        conversion = status_counts["Согласовано"] / total_with_status * 100
                        st.progress(min(conversion / 100, 1.0), text=f"Конверсия: {conversion:.1f}%")
                
                if st.session_state.user_info and st.session_state.user_info["role"] == "admin" and "_user" in df.columns:
                    st.divider()
                    st.markdown("### 👥 По менеджерам")
                    user_stats = df.groupby("_user").agg(
                        Количество=("итоговая_стоимость", "count"),
                        Сумма=("итоговая_стоимость", "sum"),
                        Средний=("итоговая_стоимость", "mean")
                    ).reset_index()
                    st.dataframe(user_stats, use_container_width=True, hide_index=True)
                
                st.divider()
                c1, c2 = st.columns(2)
                with c1:
                    if "тип_объекта" in df.columns:
                        tc = df["тип_объекта"].value_counts().reset_index()
                        tc.columns = ["Тип", "Кол-во"]
                        fig1 = px.pie(tc, values="Кол-во", names="Тип", hole=0.4, title="По типам объектов")
                        st.plotly_chart(fig1, use_container_width=True)
                with c2:
                    if "тип_объекта" in df.columns:
                        dg = df.groupby("тип_объекта")["итоговая_стоимость"].sum().reset_index()
                        fig2 = px.bar(dg, x="тип_объекта", y="итоговая_стоимость", color="тип_объекта", title="Стоимость по типам")
                        fig2.update_layout(showlegend=False)
                        st.plotly_chart(fig2, use_container_width=True)
                
                if "_время" in df.columns and len(df) >= 2:
                    st.divider()
                    st.markdown("### 📈 Динамика по времени")
                    df_time = df.copy()
                    df_time["дата"] = pd.to_datetime(df_time["_время"]).dt.date
                    daily = df_time.groupby("дата").agg(
                        Количество=("итоговая_стоимость", "count"),
                        Сумма=("итоговая_стоимость", "sum")
                    ).reset_index()
                    fig3 = px.line(daily, x="дата", y="Сумма", title="Сумма КП по дням", markers=True)
                    st.plotly_chart(fig3, use_container_width=True)
                
                st.divider()
                st.dataframe(df[["_user", "_время", "тип_объекта", "площадь_квм", "итоговая_стоимость"]].tail(10), use_container_width=True, hide_index=True)
                
                if st.session_state.user_info and st.session_state.user_info["role"] == "admin":
                    if st.button("🗑️ Очистить ВСЮ историю", use_container_width=True):
                        clear_all_history()
                        st.rerun()
                else:
                    if st.button("🗑️ Очистить мою историю", use_container_width=True):
                        clear_user_history()
                        st.rerun()
        
        with subtab2:
            st.markdown("### 🔄 Сравнение двух КП")
            
            if len(st.session_state.history) < 2:
                st.warning("Нужно минимум 2 расчёта для сравнения.")
            else:
                kp_list = [f"{h.get('_время', '')} | {h.get('тип_объекта', '')} | {h.get('итоговая_стоимость', 0):,.0f} ₽".replace(",", " ") for h in st.session_state.history]
                
                col_a, col_b = st.columns(2)
                with col_a:
                    idx1 = st.selectbox("КП №1:", range(len(kp_list)), format_func=lambda x: kp_list[x])
                with col_b:
                    idx2 = st.selectbox("КП №2:", range(len(kp_list)), format_func=lambda x: kp_list[x], index=min(1, len(kp_list)-1))
                
                if idx1 != idx2:
                    kp1 = st.session_state.history[idx1]
                    kp2 = st.session_state.history[idx2]
                    
                    comp_data = {
                        "Параметр": ["Тип объекта", "Площадь, м²", "Базовая ставка, ₽/м²", "Базовая стоимость, ₽", "После коэффициентов, ₽", "Итоговая стоимость, ₽", "Срок, мес.", "Менеджер"],
                        "КП №1": [kp1.get("тип_объекта", ""), f"{kp1.get('площадь_квм', 0):,}".replace(",", " "), f"{kp1.get('базовая_ставка_за_квм', 0):,}".replace(",", " "), f"{kp1.get('базовая_стоимость', 0):,.0f}".replace(",", " "), f"{kp1.get('стоимость_после_коэффициентов', 0):,.0f}".replace(",", " "), f"{kp1.get('итоговая_стоимость', 0):,.0f}".replace(",", " "), str(kp1.get("срок_месяцев", "")), kp1.get("_user", "")],
                        "КП №2": [kp2.get("тип_объекта", ""), f"{kp2.get('площадь_квм', 0):,}".replace(",", " "), f"{kp2.get('базовая_ставка_за_квм', 0):,}".replace(",", " "), f"{kp2.get('базовая_стоимость', 0):,.0f}".replace(",", " "), f"{kp2.get('стоимость_после_коэффициентов', 0):,.0f}".replace(",", " "), f"{kp2.get('итоговая_стоимость', 0):,.0f}".replace(",", " "), str(kp2.get("срок_месяцев", "")), kp2.get("_user", "")]
                    }
                    
                    comp_df = pd.DataFrame(comp_data)
                    
                    diff = kp2.get('итоговая_стоимость', 0) - kp1.get('итоговая_стоимость', 0)
                    diff_pct = (diff / kp1.get('итоговая_стоимость', 1)) * 100 if kp1.get('итоговая_стоимость', 0) else 0
                    
                    if diff > 0:
                        st.warning(f"💰 КП №2 дороже на {diff:,.0f} ₽ ({diff_pct:+.1f}%)".replace(",", " "))
                    elif diff < 0:
                        st.success(f"💰 КП №2 дешевле на {abs(diff):,.0f} ₽ ({diff_pct:+.1f}%)".replace(",", " "))
                    else:
                        st.info("💰 Стоимость одинакова")
                    
                    st.dataframe(comp_df, use_container_width=True, hide_index=True)
                else:
                    st.warning("Выберите два разных КП для сравнения.")

# ========== ВКЛАДКА 3: ЧАТ С ИИ ==========
with tab_chat:
    st.title("💬 Чат с ИИ-консультантом")
    st.markdown("Задайте вопрос по расчётам, ценам или условиям.")
    
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    
    chat_input = st.text_area("Ваш вопрос:", placeholder="Какая скидка за предоплату?", height=80, key="chat_input")
    
    col_c1, col_c2, col_c3 = st.columns([1, 1, 3])
    with col_c1:
        if st.button("📨 Отправить", use_container_width=True, type="primary"):
            if chat_input:
                with st.spinner("🤔 ИИ думает..."):
                    chat_prompt = f"""
Ты — ИИ-консультант строительной компании. Отвечай кратко, используя базу знаний.

=== БАЗА ЗНАНИЙ ===
{st.session_state.price_list}
=== КОНЕЦ БАЗЫ ===

ВОПРОС: {chat_input}
"""
                    load_dotenv()
                    chat_llm = ChatOpenAI(model="deepseek/deepseek-chat", temperature=0.3, openai_api_key=os.getenv("VSEGPT_API_KEY"), openai_api_base="https://api.vsegpt.ru/v1")
                    answer = chat_llm.invoke(chat_prompt).content
                    st.session_state.chat_history.append({"user": chat_input, "bot": answer})
                    st.rerun()
    with col_c2:
        if st.button("🗑️ Очистить чат", use_container_width=True):
            st.session_state.chat_history = []
            st.rerun()
    
    if st.session_state.chat_history:
        st.divider()
        for msg in st.session_state.chat_history:
            with st.chat_message("user"):
                st.write(msg["user"])
            with st.chat_message("assistant"):
                st.write(msg["bot"])
    else:
        st.info("Задайте вопрос, чтобы начать диалог.")